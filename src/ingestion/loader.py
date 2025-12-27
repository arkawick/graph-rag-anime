"""Document loaders for various file formats."""

import os
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass
import pypdf
import docx
import markdown
from rich.console import Console

console = Console()


@dataclass
class Document:
    """Document data structure."""
    content: str
    metadata: Dict[str, Any]

    def __post_init__(self):
        """Add default metadata."""
        if "source" not in self.metadata:
            self.metadata["source"] = "unknown"
        if "type" not in self.metadata:
            self.metadata["type"] = "text"


class DocumentLoader:
    """Load documents from various file formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}

    def __init__(self):
        self.documents: List[Document] = []

    def load_file(self, file_path: str | Path) -> Document:
        """Load a single file."""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        # Load based on file type
        if extension == ".pdf":
            content = self._load_pdf(file_path)
        elif extension == ".txt":
            content = self._load_txt(file_path)
        elif extension == ".md":
            content = self._load_markdown(file_path)
        elif extension == ".docx":
            content = self._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "type": extension[1:],  # Remove the dot
            "size": file_path.stat().st_size,
        }

        return Document(content=content, metadata=metadata)

    def load_directory(
        self,
        directory: str | Path,
        recursive: bool = True,
        show_progress: bool = True
    ) -> List[Document]:
        """Load all supported documents from a directory."""
        directory = Path(directory)

        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")

        # Find all supported files
        pattern = "**/*" if recursive else "*"
        all_files = []

        for ext in self.SUPPORTED_EXTENSIONS:
            all_files.extend(directory.glob(f"{pattern}{ext}"))

        if not all_files:
            console.print(f"[yellow]No supported documents found in {directory}[/yellow]")
            return []

        # Load files
        documents = []

        if show_progress:
            from tqdm import tqdm
            file_iter = tqdm(all_files, desc="Loading documents")
        else:
            file_iter = all_files

        for file_path in file_iter:
            try:
                doc = self.load_file(file_path)
                documents.append(doc)
            except Exception as e:
                console.print(f"[red]Error loading {file_path}: {e}[/red]")

        self.documents.extend(documents)

        if show_progress:
            console.print(f"[green]✓ Loaded {len(documents)} documents[/green]")

        return documents

    def _load_pdf(self, file_path: Path) -> str:
        """Load PDF file."""
        reader = pypdf.PdfReader(str(file_path))
        text = ""

        for page in reader.pages:
            text += page.extract_text() + "\n\n"

        return text.strip()

    def _load_txt(self, file_path: Path) -> str:
        """Load text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _load_markdown(self, file_path: Path) -> str:
        """Load markdown file."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # Convert markdown to plain text (strip HTML tags)
        html = markdown.markdown(content)
        # Simple HTML tag removal (for better indexing)
        import re
        text = re.sub(r'<[^>]+>', '', html)

        return text

    def _load_docx(self, file_path: Path) -> str:
        """Load DOCX file."""
        doc = docx.Document(str(file_path))
        text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])

        return text.strip()
